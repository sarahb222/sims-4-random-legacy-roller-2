# -*- coding: utf-8 -*-
"""
Created on Mon Nov 25 14:51:17 2019

@author: Sarah
"""

import random
from docx import Document
from docx.shared import Pt
from openpyxl import load_workbook
import sys
import numpy as py

def randomNum(minnum, num):
    return random.randint(minnum,num)

def addPara(doc, ws, num, x):
    #Gender Law 1-3
    p = doc.add_paragraph()
    p.style.font.bold = True
    p.add_run(ws.title + ": ").font.size = Pt(16)
    p.style.font.size = Pt(12)
    p.add_run(ws.cell(row = num[x], column = 3).value).bold = False
    
def addCareer(doc, ws, val):
    p = doc.add_paragraph()
    p.style.font.size = Pt(12)
    p.add_run(ws.cell(row = val, column = 3).value).bold = False
    
def addMore(doc, ws, val, col):
    while(ws.cell(row = val, column = col).value != None):
        p = doc.add_paragraph()
        p.style.font.size = Pt(12)
        p.add_run(ws.cell(row = val, column = col).value).bold = False
        col += 1

try:
    #xl = pd.ExcelFile('FLActive.xlsx')
    global wb
    wb = load_workbook('SimsRandomWorksheet.xlsx')
    # Store configuration file values
except FileNotFoundError:
    print("ERROR")
    sys.exit()

ws_gen = wb["Gender Law"]
ws_suc = wb["Heir Law"]
ws_ms = wb["Marital Structure"]
ws_ch = wb["Children"]
ws_pc= wb["Primary Career"]
ws_sc = wb["Secondary Career"]
ws_cc = wb["Conventional Careers"]
ws_uc = wb["Unconventional Careers"]
ws_ptj = wb["Part Time Jobs"]
ws_gg = wb["Generation Goals"]
ws_mf = wb["Misc Fun"]
ws_bt = wb["Bonus Text"]
ws_bl = wb["Blood Law"]
ws_sl = wb["Species Law"]

#checking vars
check_cc = False
check_hop = False
check_child = False
check_cc_nonpro = False


num = py.zeros(9).astype(int)

num[0] = randomNum(1, ws_gen['A1'].value)
num[1] = randomNum(1, ws_suc['A1'].value)
num[2] = randomNum(1, ws_ms['A1'].value)
num[3] = randomNum(1, ws_ch['A1'].value)
num[4] = randomNum(1, ws_pc['A1'].value)
num[5] = randomNum(1, ws_gg['A1'].value)
num[6] = randomNum(1, ws_mf['A1'].value)
num[7] = randomNum(1, ws_bl['A1'].value)
num[8] = randomNum(1, ws_sl['A1'].value)

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'

addPara(doc, ws_gen, num, 0)
addPara(doc, ws_suc, num, 1)
addPara(doc, ws_bl, num, 7)
addPara(doc, ws_sl, num, 8)
addPara(doc, ws_ms, num, 2)
addPara(doc, ws_ch, num, 3)

#number of children
if num[3] <= 2:
    check_child = True

addPara(doc, ws_pc, num, 4)

#primary career
if num[4] <= 6:
    pc_val = randomNum(1, ws_cc['A1'].value)
    addCareer(doc, ws_cc, pc_val)
    check_cc = True
    if(ws_cc.cell(row = pc_val, column = 2).value == None):
        check_cc_nonpro = True
elif num[4] <= 9:
    pc_val = randomNum(1, ws_uc['A1'].value)
    addCareer(doc, ws_uc, pc_val)
else:
    pc_hop = py.zeros(5).astype(int)
    check_hop = True
    for i in range(len(pc_hop)): #edit this to remove dupes
        pc_hop[i] = randomNum(1, ws_cc['A1'].value)
    
    for i in range(len(pc_hop)):
        addCareer(doc, ws_cc, pc_hop[i])

#determine amount of secondary careers to roll
if (num[2] >=7 and num[2] <= 15) or (num[2] == 17) or (num[2] == 19) or (num[2] == 20):
    total_sc = 1
elif (num[2] == 18) or (num[2] >= 21 and num[2] <= 24):
    total_sc = 2
elif num[2] == 25:
    total_sc = 3
else:
    total_sc = 0
    
if total_sc != 0:
    num_sc = py.zeros(total_sc).astype(int)
    for i in range(len(num_sc)):
        num_sc[i] = randomNum(1, ws_sc['A1'].value)
    
    for i in range(len(num_sc)):
        addPara(doc, ws_sc, num_sc, i)
        if num_sc[i] <= 3:
            sc_val = randomNum(1, ws_cc['A1'].value)
            addCareer(doc, ws_cc, sc_val)
            check_cc = True
            if(ws_cc.cell(row = sc_val, column = 2).value == None):
                check_cc_nonpro = True
        elif num_sc[i] <= 6:
            sc_val = randomNum(1, ws_uc['A1'].value)
            addCareer(doc, ws_uc, sc_val)
        elif num_sc[i] <= 8:
            sc_val = randomNum(1, ws_ptj['A1'].value)
            addCareer(doc, ws_ptj, sc_val)
        elif num_sc[i] == 9:
            check_hop = True
            sc_hop = py.zeros(5).astype(int)
            for x in range(len(sc_hop)): #edit this to remove dupes
                sc_hop[x] = randomNum(1, ws_cc['A1'].value)
                
            for x in range(len(sc_hop)):
                addCareer(doc, ws_cc, sc_hop[x])
        else:
            break
    
#gen goal

while True:
    if(num[5] == 1 and (check_hop == True or check_cc == False)):
        num[5] = randomNum(1, ws_gg['A1'].value)
    elif(num[5] == 4 and check_cc_nonpro == False):
        num[5] = randomNum(1, ws_gg['A1'].value)
    else:
        break

addPara(doc, ws_gg, num, 5)

test_val = ws_gg.cell(row = num[5], column = 2).value
if(test_val != None):
    if(test_val == 0):
        addMore(doc, ws_gg, num[5], 4)
    else:
        bonus_ran = randomNum(ws_gg.cell(row = num[5], column = 2).value, ws_bt.cell(row = test_val, column = 2).value)
        
        #special rules for midlife crisis
        if(num[5] == 12):
            bonus_ran2 = randomNum(ws_gg.cell(row = num[5], column = 2).value, ws_bt.cell(row = test_val, column = 2).value-1)
            if(bonus_ran2 >= bonus_ran):
                bonus_ran2 += 1
            bonus_ran3 = randomNum(ws_gg.cell(row = num[5], column = 2).value, ws_bt.cell(row = test_val, column = 2).value-2)
            if(bonus_ran3 >= bonus_ran):
                bonus_ran3 += 1
                if(bonus_ran3 >= bonus_ran2):
                    bonus_ran3 += 1
            elif(bonus_ran3 >= bonus_ran2):
                bonus_ran3 += 1
                if(bonus_ran3 >= bonus_ran):
                    bonus_ran3 += 1
                    
            #all 3 numbers chosen and different, bonus_ran and 2 and 3
            #if 1, reroll misc fun
            #if 3, reroll career 
            addCareer(doc, ws_bt, bonus_ran2)
            addCareer(doc, ws_bt, bonus_ran3)
        addCareer(doc, ws_bt, bonus_ran)
        

#misc fun
while True:
    if(num[6] == 9 and check_child == True):
        num[6] = randomNum(1, ws_mf['A1'].value)
    else:
        break
        
        
addPara(doc, ws_mf, num, 6)

test_val = ws_mf.cell(row = num[6], column = 2).value
if(test_val != None):
    if(test_val == 0):
        addMore(doc, ws_mf, num[6], 4)
    else:
        bonus_ran = randomNum(ws_mf.cell(row = num[6], column = 2).value, ws_bt.cell(row = test_val, column = 2).value)
        addCareer(doc, ws_bt, bonus_ran)

doc.save("randomroll.docx")

wb.close()